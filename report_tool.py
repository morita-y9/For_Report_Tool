import pdfplumber

from docx import Document
import spacy
import streamlit as st
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

def extract_pdf_elements(pdf_path):
    elements = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=1)
            tables = page.extract_tables()
            images = page.images
            elements.append({
                "text": text,
                "tables": tables,
                "images": images
            })
    return elements

def write_to_word(elements, output_path):
    doc = Document("report_format.docx")
    base_font = "MS 明朝"

    for page_num, content in enumerate(elements):
        if content['text']:
            for line in content['text'].split('\n'):
                p = doc.add_paragraph()
                run = p.add_run(line)

                # 基本フォントスタイル
                run.font.name = base_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font)
                run.font.size = Pt(10.5)

                # 見出しと判断される行なら大きくする
                if line.strip().startswith(("1.", "2.", "3.","4.","5.","6.", "実験目的", "実験結果", "考察")):
                    run.font.size = Pt(14)
                    run.bold = True

        for table in content['tables']:
            rows, cols = len(table), len(table[0])
            t = doc.add_table(rows=rows, cols=cols)
            for i, row in enumerate(table):
                for j, cell in enumerate(row):
                    t.cell(i, j).text = cell or ''
                    for paragraph in t.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.font.name = base_font
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font)
                            run.font.size = Pt(12)

    doc.save(output_path)
    
nlp = spacy.load("en_core_web_sm")


st.title("word用実験レポート作成ツール ")
st.write("仮にレポートが吹き飛んだとしても一切の責任を負いかねます")
uploaded_file = st.file_uploader("PDFファイルをアップロード", type=["pdf"])

if uploaded_file:
    st.success("PDFファイルが正常に読み込まれました！")
    elements = extract_pdf_elements(uploaded_file)

    st.write("抽出されたページ数:", len(elements))
    
    if st.button("Wordに変換"):
        output_path = "output.docx"
        write_to_word(elements, output_path)
        with open(output_path, "rb") as f:
            st.download_button("Wordファイルをダウンロードする", f, file_name="report_draft.docx")

